"""
Generate a branded .docx Platform Overview document for Rent-A-Vacation.
Reuses brand helpers from generate_docx.py.
"""

import os
import sys
from datetime import datetime

# Add exports dir to path so we can import helpers
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

from generate_docx import (
    create_branded_doc,
    add_logo_header,
    add_page_numbers,
    add_metadata,
    add_body,
    add_blockquote,
    add_table_from_data,
    add_horizontal_rule,
    add_footer,
    DEEP_TEAL,
    WARM_CORAL,
    DARK_NAVY,
    BRAND_FONT,
)
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_numbered_list(doc, items, bold_prefix=False):
    """Add a numbered list with brand styling."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{i}. ")
        run.font.name = BRAND_FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DEEP_TEAL

        # Check for "bold → rest" pattern
        if "\u2192" in item:
            parts = item.split("\u2192", 1)
            run2 = p.add_run(parts[0].strip())
            run2.font.name = BRAND_FONT
            run2.font.size = Pt(10)
            run2.font.bold = True
            run2.font.color.rgb = DARK_NAVY
            run3 = p.add_run(f" \u2192 {parts[1].strip()}")
            run3.font.name = BRAND_FONT
            run3.font.size = Pt(10)
            run3.font.color.rgb = DARK_NAVY
        else:
            run2 = p.add_run(item)
            run2.font.name = BRAND_FONT
            run2.font.size = Pt(10)
            run2.font.color.rgb = DARK_NAVY


def add_bullet_list(doc, items):
    """Add a bullet list with brand styling."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

        # Check for "Label: Description" pattern
        if ": " in item:
            label, desc = item.split(": ", 1)
            run_bullet = p.add_run("\u2022 ")
            run_bullet.font.name = BRAND_FONT
            run_bullet.font.size = Pt(10)
            run_bullet.font.color.rgb = DEEP_TEAL
            run_label = p.add_run(f"{label}: ")
            run_label.font.name = BRAND_FONT
            run_label.font.size = Pt(10)
            run_label.font.bold = True
            run_label.font.color.rgb = DARK_NAVY
            run_desc = p.add_run(desc)
            run_desc.font.name = BRAND_FONT
            run_desc.font.size = Pt(10)
            run_desc.font.color.rgb = DARK_NAVY
        else:
            run = p.add_run(f"\u2022 {item}")
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_NAVY


def generate_platform_overview():
    today = datetime.now().strftime("%m%d%Y")
    doc = create_branded_doc("Platform Overview")
    add_logo_header(doc, doc_title="Platform Overview \u2014 What\u2019s Been Built")
    add_page_numbers(doc)

    # Metadata
    add_metadata(doc, [
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Version", "v0.9.0 (Pre-Launch)"),
        ("Website", "https://rent-a-vacation.com"),
        ("Repository", "github.com/rent-a-vacation/rav-website"),
    ])

    # ── What It Is ──
    doc.add_heading("What It Is", level=1)
    add_body(
        doc,
        "An AI-powered marketplace where vacation club owners rent directly to travelers, "
        "and travelers name their price. RAV earns a 15% default commission "
        "(Pro \u22122%, Business \u22125%) on successful bookings. Think Airbnb, but purpose-built "
        "for timeshare inventory across Hilton, Marriott, Disney, and 6 other vacation "
        "club brands (117 resorts, 351 unit types). Two marketplace primitives no "
        "competitor offers: Name Your Price (bidding on any listing) and RAV Wishes "
        "(travelers post dream trips; verified owners compete with proposals).",
    )
    add_body(
        doc,
        "Tagline: \u201cName Your Price. Book Your Paradise.\u201d",
    )

    # ── Tech Stack ──
    doc.add_heading("Tech Stack", level=1)
    add_bullet_list(doc, [
        "Frontend: React 18 + TypeScript + Vite + Tailwind + shadcn/ui",
        "Backend: Supabase (PostgreSQL, Auth, Edge Functions, RLS)",
        "Payments: Stripe (checkout, Connect payouts, webhooks, Tax-ready)",
        "Voice: VAPI (Deepgram Nova-3 STT + GPT-4o-mini + ElevenLabs TTS)",
        "Text Chat: OpenRouter (RAVIO assistant, SSE streaming)",
        "Notifications: Resend (email) + Twilio (SMS, TCPA-compliant)",
        "Observability: Sentry (errors + session replay) + GA4 (consent-gated)",
        "Deployment: Vercel (frontend) + Supabase (backend)",
    ])

    add_horizontal_rule(doc)

    # ── Core User Journeys ──
    doc.add_heading("Core User Journeys", level=1)

    doc.add_heading("Property Owner Flow", level=2)
    add_numbered_list(doc, [
        "Sign up \u2192 pending approval by RAV admin",
        "Add property (9 brands supported) \u2192 create listing with nightly rate",
        "Listing goes to pending_approval \u2192 RAV admin approves/rejects",
        "Once booked \u2192 owner confirms resort reservation \u2192 RAV verifies \u2192 escrow holds funds",
        "After checkout + 5 days \u2192 funds released \u2192 Stripe Connect payout",
    ])

    doc.add_heading("Traveler Flow", level=2)
    add_numbered_list(doc, [
        "Browse/search listings (voice search, text chat, filters)",
        "View property details with fair value scoring",
        "Place bids or propose alternate dates",
        "Checkout via Stripe \u2192 booking confirmed",
        "Track booking in My Bookings, file disputes if needed",
    ])

    doc.add_heading("Admin Flow (RAV Ops)", level=2)
    add_numbered_list(doc, [
        "Unified operations dashboard: Users, Listings, Bookings, Escrow, Payouts, Financials, Disputes, Voice, Resorts, API Keys, Notifications, Events",
        "Approve/reject listings and users with bulk actions + audit trail",
        "Manage escrow lifecycle (verify, hold, release, refund)",
        "Dispute resolution queue with assignment + evidence review",
        "Voice search monitoring, tier/user quota overrides, usage dashboard",
        "Resort CSV import with validation + duplicate detection",
        "Notification catalog, seasonal event calendar, delivery log",
    ])

    doc.add_heading("Executive Flow (RAV Insights)", level=2)
    add_numbered_list(doc, [
        "6-section BI dashboard at /executive-dashboard",
        "Headline KPIs, GMV trend, revenue waterfall, bid activity",
        "Marketplace health: Liquidity Score, supply/demand map",
        "Market intelligence with BYOK pattern (AirDNA, STR)",
        "Industry feed + regulatory radar + unit economics",
    ])

    add_horizontal_rule(doc)

    # ── Features Built ──
    doc.add_heading("Features Built Across 48+ Sessions", level=1)

    features_data = [
        ("Auth", "Email/password + Google OAuth, role-based access (6 roles), email verification, user approval workflow, realtime role upgrade"),
        ("Listings", "Create/edit listings, nightly pricing, fair value scoring, photo uploads, dynamic pricing (urgency/seasonal/demand), admin edit with audit trail"),
        ("Name Your Price (Bidding)", "Bid on any listing, propose alternate dates, 24hr expiry, owner accept/reject/counter, Bid Spread Index"),
        ("RAV Wishes (Reverse Auction)", "Travelers post dream trips; verified owners send proposals; traveler picks winner. Formerly \u201cVacation Wishes\u201d (Session 47 rebrand)"),
        ("RAV Deals", "Curated distressed/expiring-week inventory surface, feeds directly into bidding"),
        ("Booking", "Stripe Checkout, fee breakdown (base + service + cleaning + tax), 5-step booking timeline, cancellation policy UI"),
        ("Payments", "Stripe Connect (owner onboarding + payouts), webhooks (6 events), escrow lifecycle, Stripe Tax-ready (awaiting #127)"),
        ("Cancellation", "Policy-based (flexible/moderate/strict/super_strict) renter cancellation, owner cancellation with full refund, automated Stripe refunds"),
        ("Escrow (PaySafe)", "6-status lifecycle, owner confirmation timer, RAV verification, auto-release after checkout+5d, hold/unhold, refund"),
        ("Disputes", "Expanded dispute categories (renter + owner), evidence upload, admin queue with assignment + resolution"),
        ("Voice Search (Ask RAVIO)", "VAPI integration, Deepgram Nova-3, tier-based quotas, admin overrides, usage dashboard, observability + alert thresholds"),
        ("Text Chat (Chat with RAVIO)", "Streaming AI assistant via OpenRouter with property cards"),
        ("RAV Smart Suite", "5 free public tools at /tools: SmartEarn, SmartPrice, SmartCompare, SmartMatch, SmartBudget"),
        ("TrustShield", "Multi-step owner identity + ownership verification; trust-level progression"),
        ("ResortIQ", "117 resorts, 351 unit types, 9 brands; auto-populates listing forms, powers voice search"),
        ("Renter Dashboard (My Trips)", "/my-trips \u2014 4 tabs (Overview, Bookings, Offers, Favorites) with check-in countdown"),
        ("My Rentals (Owner Dashboard)", "Consolidated 11 tabs \u2192 4: Dashboard, My Listings, Bookings & Earnings, Account. Formerly \u201cOwner\u2019s Edge\u201d"),
        ("RAV Ops (Admin)", "Unified operations: Users, Listings, Bookings, Escrow, Payouts, Financials, Disputes, Voice, Resorts, API Keys, Notifications, Events"),
        ("RAV Insights (Executive BI)", "6-section dashboard: KPIs, GMV, marketplace health, market intel (BYOK), industry feed, unit economics. Formerly \u201cRAV Command\u201d"),
        ("Notification Center", "Multi-channel routing (in-app/email/SMS), per-type preferences, TCPA opt-in, seasonal events, delivery log"),
        ("SMS Infrastructure", "Twilio integration: notification-dispatcher, sms-scheduler, twilio-webhook. SMS_TEST_MODE=true \u2014 prod gated on A2P 10DLC (#127)"),
        ("Public API + Developer Portal", "OpenAPI 3.0 spec, /developers Swagger UI, API key management, tiered rate limits, IP allowlisting (CIDR)"),
        ("Referral Program", "Unique codes, tracking dashboard, signup attribution via ?ref=CODE"),
        ("Pre-Booking Messaging", "Inquiry threads between travelers and owners before booking"),
        ("Saved Searches", "Save criteria + price drop alerts"),
        ("Destinations Explorer", "10 destinations, 35 cities \u2014 curated discovery pages"),
        ("Compare Properties", "Up to 3 listings side-by-side with best-value badges"),
        ("iCal Export", "RFC 5545 calendar export for owner bookings"),
        ("Realtime", "useRealtimeSubscription replaced all polling (notifications, messages, unread counts)"),
        ("SEO", "Meta tags, sitemap, robots.txt, FAQ JSON-LD, OG images, JSON-LD on /tools"),
        ("Security", "CSP headers, rate limiting, RLS on every table, IP allowlisting for API keys"),
        ("GDPR", "Data export, account deletion with 14-day grace period, cookie consent"),
        ("Architecture Docs", "Auto-generated flow diagrams from declarative Flow Manifests at /architecture"),
        ("PWA", "Service worker, installable, offline-capable"),
        ("Observability", "Sentry (errors + session replay) + GA4 (consent-gated, G-G2YCVHNS25)"),
    ]

    add_table_from_data(
        doc,
        ["Area", "What\u2019s Built"],
        features_data,
    )

    add_horizontal_rule(doc)

    # ── Current Numbers ──
    doc.add_heading("Current Numbers", level=1)

    add_table_from_data(
        doc,
        ["Metric", "Count"],
        [
            ("Automated tests", "956 passing (121 test files)"),
            ("P0 critical-path tests", "97 (`@p0` tagged, ~2s run)"),
            ("Database migrations", "046 (all deployed to DEV + PROD)"),
            ("Edge functions", "30 (27 in PROD; 3 SMS functions DEV-only until A2P 10DLC)"),
            ("Resorts / unit types", "117 / 351 across 9 brands"),
            ("Routes (pages)", "~35 (incl. /developers, /tools, /destinations, /notifications)"),
            ("Type errors / Lint errors", "0 / 0"),
            ("Build status", "Clean"),
            ("Sessions shipped", "48+ in 19 months"),
            ("Platform uptime (staging)", "99.97%"),
            ("dev vs main", "In sync (PR #239 merged Apr 1)"),
        ],
    )

    add_horizontal_rule(doc)

    # ── Remaining Pre-Launch Items ──
    doc.add_heading("Remaining Pre-Launch Items", level=1)
    add_body(
        doc,
        "The platform is feature-complete. Remaining blockers are external (legal, "
        "business formation, and third-party verification) rather than engineering work.",
    )

    add_table_from_data(
        doc,
        ["#", "Issue", "Status"],
        [
            ("#127", "Business Formation (LLC + EIN)", "BLOCKER \u2014 gates Stripe Tax, Puzzle.io, Mercury bank, A2P 10DLC SMS"),
            ("#80", "Legal review: ToS and Privacy Policy", "Awaiting legal counsel review"),
            ("#187", "Manual verification workflow (TrustShield)", "Operational process to document"),
            ("#230-234", "Social media account setup", "Facebook, Instagram, LinkedIn, Twitter, GBP"),
            ("#87", "Launch readiness checklist", "Ready when #127 and #80 close"),
        ],
    )
    add_body(
        doc,
        "Already shipped (previously on this list): GA4 (#74), Admin Tax Reporting (#62), "
        "1099-K Compliance (#64), Stripe Connect webhooks, Sentry, CSP/rate-limits, GDPR.",
    )

    add_horizontal_rule(doc)

    # ── Current Platform State ──
    doc.add_heading("Current Platform State", level=1)
    add_bullet_list(doc, [
        "PROD: Staff Only Mode enabled \u2014 platform locked for pre-launch control",
        "Stripe Tax: Code ready (automatic_tax enabled), not activated pending LLC/EIN (#127)",
        "SMS: Deployed to DEV with SMS_TEST_MODE=true; production traffic gated on A2P 10DLC",
        "Puzzle.io accounting: Onboarding paused at step 7 pending LLC/EIN (#127)",
        "Repository: Private (changed Mar 4, 2026) \u2014 branch protection requires GitHub Team plan",
        "Supabase CLI: Currently linked to DEV project",
        "dev and main: In sync (PR #239 merged Apr 1, 2026)",
    ])

    add_horizontal_rule(doc)

    # ── Brand Architecture ──
    doc.add_heading("Brand Architecture (Session 47 Rebrand \u2014 Apr 12, 2026)", level=1)
    add_body(
        doc,
        "The RAV brand family follows three naming patterns: RAV-prefix for platform "
        "features and internal tools; CompoundName for the trust/infrastructure layer; "
        "plain language for customer-facing navigation.",
    )
    add_table_from_data(
        doc,
        ["Canonical Name", "Type", "Notes"],
        [
            ("Name Your Price", "Primary \u2014 bidding", "Hero feature, master tagline"),
            ("RAV Wishes", "Primary \u2014 reverse auction", "Formerly \u201cVacation Wishes\u201d"),
            ("RAV Deals", "Primary \u2014 discovery", "New distressed-inventory surface"),
            ("TrustShield + PaySafe", "Supporting \u2014 trust", "Verification + escrow"),
            ("Ask RAVIO / Chat with RAVIO", "Supporting \u2014 AI", "Voice + text search"),
            ("RAV Smart[X]", "Supporting \u2014 tools", "SmartEarn, SmartPrice, SmartCompare, SmartMatch, SmartBudget"),
            ("My Trips / My Rentals", "Nav \u2014 customer", "Plain-language dashboards"),
            ("RAV Insights", "Infrastructure \u2014 BI", "Formerly \u201cRAV Command\u201d"),
            ("RAV Ops", "Infrastructure \u2014 admin", "Formerly \u201cAdmin Dashboard\u201d"),
            ("ResortIQ", "Infrastructure \u2014 data", "117 resorts, 351 unit types"),
        ],
    )

    # Footer
    add_footer(doc, "Rent-A-Vacation \u2022 Confidential \u2022 Generated " + datetime.now().strftime("%B %d, %Y"))

    # Save
    output_path = os.path.join(SCRIPT_DIR, f"RAV-Platform-Overview-{today}.docx")
    doc.save(output_path)
    print(f"Generated: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_platform_overview()
